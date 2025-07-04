# coding=utf-8
from docx import Document
from dateutil import parser
from itertools import groupby
import os, locale, re

def read_weekday(filename):
    wordDoc = Document(os.getcwd() + f"/{filename}.docx")
    #заменяем название в файле:
    if filename.find('января') != -1:
        filename = filename.replace('января', 'January')
    elif filename.find('февраля') != -1:
        filename = filename.replace('февраля', 'February')
    elif filename.find('марта') != -1:
        filename = filename.replace('марта', 'March')
    elif filename.find('апреля') != -1:
        filename = filename.replace('апреля', 'April')
    elif filename.find('мая') != -1:
        filename = filename.replace('мая', 'May')
    elif filename.find('июня') != -1:
        filename = filename.replace('июня', 'June')
    elif filename.find('июля') != -1:
        filename = filename.replace('июля', 'July')
    elif filename.find('августа') != -1:
        filename = filename.replace('августа', 'August')
    elif filename.find('сентября') != -1:
        filename = filename.replace('сентября', 'September')
    elif filename.find('октября') != -1:
        filename = filename.replace('октября', 'October')
    elif filename.find('ноября') != -1:
        filename = filename.replace('ноября', 'November')
    elif filename.find('декабря') != -1:
        filename = filename.replace('декабря', 'December')

    weekday_list = {"Понедельник":1,"Вторник":2,"Среда":3,"Четверг":4,"Пятница":5,"Суббота":6,"Воскресенье":7}
    weekday_list1 = {value:key for key, value in weekday_list.items()}

    #получаем на какой день недели расписание
    locale.setlocale(locale.LC_TIME, 'ru_RU.utf8')

    print('Файл',filename)
    extr_date = re.search("(\d+.+?)г", filename).group(1)
    week_day = parser.parse(extr_date).weekday()+1

    #получаем день недели в виде букв
    liva_weekday = weekday_list1.get(week_day)
    print('Есть замены на: ', liva_weekday)
    return liva_weekday

def read_file(group, filename, merge):
    wordDoc = Document(os.getcwd() + f"/{filename}.docx")
    #заменяем название в файле:
    if filename.find('января') != -1:
        filename = filename.replace('января', 'January')
    elif filename.find('февраля') != -1:
        filename = filename.replace('февраля', 'February')
    elif filename.find('марта') != -1:
        filename = filename.replace('марта', 'March')
    elif filename.find('апреля') != -1:
        filename = filename.replace('апреля', 'April')
    elif filename.find('мая') != -1:
        filename = filename.replace('мая', 'May')
    elif filename.find('июня') != -1:
        filename = filename.replace('июня', 'June')
    elif filename.find('июля') != -1:
        filename = filename.replace('июля', 'July')
    elif filename.find('августа') != -1:
        filename = filename.replace('августа', 'August')
    elif filename.find('сентября') != -1:
        filename = filename.replace('сентября', 'September')
    elif filename.find('октября') != -1:
        filename = filename.replace('октября', 'October')
    elif filename.find('ноября') != -1:
        filename = filename.replace('ноября', 'November')
    elif filename.find('декабря') != -1:
        filename = filename.replace('декабря', 'December')

    weekday_list = {"Понедельник":1,"Вторник":2,"Среда":3,"Четверг":4,"Пятница":5,"Суббота":6,"Воскресенье":7}
    weekday_list1 = {value:key for key, value in weekday_list.items()}

    #получаем на какой день недели расписание
    locale.setlocale(locale.LC_TIME, 'ru_RU.utf8')

    extr_date = re.search("(\d+.+?)г", filename).group(1)
    week_day = parser.parse(extr_date).weekday()+1

    #получаем день недели в виде букв
    liva_weekday = weekday_list1.get(week_day)
    print('Будет показаны замены на: ', liva_weekday)

    Substitutions = []
    Substitutions2 = []
    name = []
    nameNEW = []
    time = []
    result = []
    for table in wordDoc.tables:
        for row in table.rows:
            if group.upper() in row.cells[0].text:
                s = row.cells[2].text
                s2 = row.cells[4].text
                print(s, s2)
                patter = r'[А-ЯЁа-яё]+\s[А-ЯЁа-яё]{1}\.\s*[А-ЯЁа-яё]{1}\.\s*\/*\s*[А-ЯЁа-яё]*\s*[А-ЯЁа-яё]*\.*\s*[А-ЯЁа-яё]*\.*' # ищем фамилии

                if type(re.findall(patter, s)) != type(None) or s != 'н/б' or type(re.findall(patter, s2)) != type(None) or s2 != 'н/б':
                    name.append(re.findall(patter, s)) # заносим найденные фамилии в ячейках s в список name
                    nameNEW.append(re.findall(patter, s2)) # заносим найденные фамилии в ячейках s2 в список name

                s = re.sub(r'\([^()]*\)', '', s)
                s2 = re.sub(r'\([^()]*\)', '', s2)

                if 'Занятия с ' in s2:
                    patterTIME = r'[0-9][0-9]\:[0-9][0-9]'
                    time.append(re.findall(patterTIME, s2))
                    s2 = re.sub(patterTIME, '', s2)
                    s2 = s2.replace('Занятия с ', '')

                s = eval('"' + s.replace('\n','') + '"')
                s2 = eval('"' + s2.replace('\n','') + '"')

                Substitutions2.append(row.cells[1].text + '. ' + s2 )
                Substitutions.append('\n' + row.cells[1].text + '. ' + '&#0822;' + '&#0822;'.join(s)  + s2 )




    #ФОРМАТИРУМ МАТРИЦУ С ФАМИЛИЯМИ
    name = ([x for x in name if x])
    name = [el for el, _ in groupby(name)]
    name = sum(name, [])

    name = [line.rstrip() for line in name] #удаляем символы \n из строки фамилий

    print('NAME:', name, 'TIME:', time)

    #ФОРМАТИРУМ МАТРИЦУ С ФАМИЛИЯМИ НОВЫХ ПАР
    nameNEW = ([x for x in nameNEW if x])
    nameNEW = [el for el, _ in groupby(nameNEW)]
    nameNEW = sum(nameNEW, [])

    nameNEW = [line.rstrip() for line in nameNEW] #удаляем символы \n из строки фамилий

    count = 0
    for item_i in range(len(Substitutions2)):
        Substitutions2[item_i] += '  (' + nameNEW[count] + ')' #добавляем скобки
        count += 1

    print(Substitutions2)

    if time:
        #ФОРМАТИРУМ МАТРИЦУ С ВРЕМЕНЕМ
        time = ([x for x in time if x])
        time = [el for el, _ in groupby(time)]
        time = sum(time, [])

        time = [line.rstrip() for line in time] #удаляем символы \n из строки времени

        print('TIME:', time, 'Пара начинается в', time[0], ' (изм.)\n\n')

        time = ['В колледж ' + str(time[0]) + ' (изм.)']
        result = time + Substitutions
    else:
        result = Substitutions

    if not merge:
        print('Замены:', result)
        return result
    elif merge:
        print('Замены2:', Substitutions2)
        return Substitutions2, time if time else 'Пусто'

#read_weekday('Замена на 13 сентября 2021 г.')
read_file('4исп-9', 'Замена на 15 ноября 2021 г.', True)